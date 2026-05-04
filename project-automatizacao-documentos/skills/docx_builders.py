from __future__ import annotations

from datetime import datetime, timedelta
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


def _style(doc: Document, colors: dict[str, str]) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(colors["text_color"])

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Gerado localmente pelo projeto de automação de documentos")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string("64748B")


def _shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _border(cell, color: str = "D9E2EC", size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def _set_widths(table, widths_cm: list[float]) -> None:
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(width)


def _set_cell_text(cell, text: str, size: float = 9.5, bold: bool = False, color: str = "1F2933", align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(text))
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _p(doc: Document, text: str, size: float = 10.5, bold: bool = False, color: str = "1F2933") -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.08
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def _bullets(doc: Document, items: list[str]) -> None:
    if not items:
        _p(doc, "Sem informações adicionais registradas.")
        return
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.05
        run = paragraph.add_run(item)
        run.font.name = "Arial"
        run.font.size = Pt(10)


def _heading(doc: Document, text: str, colors: dict[str, str], level: int = 1) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(10 if level == 1 else 7)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(16 if level == 1 else 12.5)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(colors["primary_color"] if level == 1 else colors["secondary_color"])


def _title_block(doc: Document, title: str, subtitle: str, project_name: str, colors: dict[str, str]) -> None:
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(2)
    r1 = title_p.add_run(title)
    r1.font.name = "Arial"
    r1.font.size = Pt(21)
    r1.bold = True
    r1.font.color.rgb = RGBColor.from_string(colors["primary_color"])

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(3)
    r2 = subtitle_p.add_run(project_name)
    r2.font.name = "Arial"
    r2.font.size = Pt(13)
    r2.bold = True
    r2.font.color.rgb = RGBColor.from_string(colors["text_color"])

    if subtitle:
        subtitle2 = doc.add_paragraph()
        subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle2.paragraph_format.space_after = Pt(10)
        r3 = subtitle2.add_run(subtitle)
        r3.font.name = "Arial"
        r3.font.size = Pt(10.5)
        r3.font.color.rgb = RGBColor.from_string(colors["secondary_color"])


def _key_value_table(doc: Document, rows: list[tuple[str, str]], colors: dict[str, str]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _set_widths(table, [4.9, 12.0])
    for idx, (key, value) in enumerate(rows):
        left = table.cell(idx, 0)
        right = table.cell(idx, 1)
        _set_cell_text(left, key, bold=True, color="FFFFFF")
        _set_cell_text(right, value)
        for cell in (left, right):
            _border(cell)
        _shade(left, colors["secondary_color"])
        if idx % 2 == 1:
            _shade(right, "F8FAFC")
    doc.add_paragraph()


def _table(
    doc: Document,
    headers: list[str],
    rows: list[list[Any]],
    widths: list[float],
    colors: dict[str, str],
    font_size: float = 9.0,
    center_columns: set[int] | None = None,
) -> None:
    center_columns = center_columns or set()
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _set_widths(table, widths)

    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        _set_cell_text(cell, header, size=font_size, bold=True, color="FFFFFF", align=WD_ALIGN_PARAGRAPH.CENTER)
        _shade(cell, colors["secondary_color"])
        _border(cell)

    for row_idx, row in enumerate(rows, start=1):
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.CENTER if idx in center_columns else WD_ALIGN_PARAGRAPH.LEFT
            _set_cell_text(cells[idx], str(value), size=font_size, align=align)
            _border(cells[idx])
        if row_idx % 2 == 0:
            for cell in cells:
                _shade(cell, "F8FAFC")
    doc.add_paragraph()


def _callout(doc: Document, title: str, body: str, colors: dict[str, str], fill: str | None = None) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _set_widths(table, [16.9])
    cell = table.cell(0, 0)
    _shade(cell, fill or colors["light_fill"])
    _border(cell, "BFD0E0", "8")
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(3)
    title_run = paragraph.add_run(title)
    title_run.font.name = "Arial"
    title_run.font.size = Pt(10.5)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor.from_string(colors["primary_color"])
    paragraph2 = cell.add_paragraph()
    paragraph2.paragraph_format.space_after = Pt(0)
    body_run = paragraph2.add_run(body)
    body_run.font.name = "Arial"
    body_run.font.size = Pt(10)
    body_run.font.color.rgb = RGBColor.from_string(colors["text_color"])
    doc.add_paragraph()


def _identity(scope_context: dict[str, Any], answers: dict[str, Any]) -> dict[str, str]:
    project = scope_context["project"]
    return {
        "project_name": answers.get("project_name") or project.get("project_name") or "Projeto",
        "client_name": answers.get("client_name") or project.get("client_name") or "",
        "contractor_name": answers.get("contractor_name") or project.get("contractor_name") or "",
        "coordinator_name": answers.get("coordinator_name") or project.get("coordinator_name") or "",
    }


def _metadata_rows(identity: dict[str, str], document_name: str, date_value: str, extra: list[tuple[str, str]] | None = None) -> list[tuple[str, str]]:
    rows = [
        ("Projeto", identity["project_name"]),
        ("Cliente", identity["client_name"]),
        ("Responsável", identity["coordinator_name"]),
        ("Documento", document_name),
        ("Data", date_value),
    ]
    if identity["contractor_name"]:
        rows.insert(2, ("Contratada", identity["contractor_name"]))
    if extra:
        rows.extend(extra)
    return rows


def _timeline_periods(start_date: str, sprint_length_days: int, phases: list[dict[str, Any]]) -> list[dict[str, str]]:
    current = datetime.strptime(start_date, "%d/%m/%Y")
    rows = []
    for phase in phases:
        duration_sprints = int(phase["duration_sprints"])
        days = duration_sprints * sprint_length_days
        end = current + timedelta(days=days - 1)
        rows.append(
            {
                "name": phase["name"],
                "objective": phase["objective"],
                "period": f"{current.strftime('%d/%m/%Y')} a {end.strftime('%d/%m/%Y')}",
            }
        )
        current = end + timedelta(days=1)
    return rows


def build_change_request(scope_context: dict[str, Any], answers: dict[str, Any], output_path: Path, runtime: dict[str, Any]) -> Path:
    colors = runtime["docx"]
    identity = _identity(scope_context, answers)
    doc = Document()
    _style(doc, colors)
    _title_block(doc, "Solicitação de Mudança de Escopo", answers["request_title"], identity["project_name"], colors)
    _key_value_table(
        doc,
        _metadata_rows(
            identity,
            "Solicitação de Mudança de Escopo",
            answers["request_date"],
            [("ID da mudança", answers["request_id"]), ("Solicitante", answers["requester_name"])],
        ),
        colors,
    )
    _heading(doc, "Descrição da solicitação", colors)
    _p(doc, answers["requested_change"])
    _heading(doc, "Justificativa", colors)
    _p(doc, answers["business_reason"])
    _heading(doc, "Impactos preliminares", colors)
    _bullets(doc, answers["impacts"])
    _heading(doc, "Recomendação", colors)
    _p(doc, answers["decision"])
    _heading(doc, "Aprovação indicativa", colors)
    _table(
        doc,
        ["Papel", "Nome", "Registro"],
        [["Solicitante", answers["requester_name"], answers["request_date"]], ["Aprovador", answers["approver_name"], ""]],
        [4.5, 7.0, 5.4],
        colors,
        center_columns={0, 2},
    )
    doc.save(output_path)
    return output_path


def build_status_report(scope_context: dict[str, Any], answers: dict[str, Any], output_path: Path, runtime: dict[str, Any]) -> Path:
    colors = runtime["docx"]
    identity = _identity(scope_context, answers)
    doc = Document()
    _style(doc, colors)
    _title_block(doc, "Status Report", answers["cycle_name"], identity["project_name"], colors)
    _key_value_table(
        doc,
        _metadata_rows(
            identity,
            "Status Report",
            answers["report_date"],
            [("Período de referência", answers["reference_period"]), ("Status geral", answers["overall_status"]), ("Identificador", answers["report_id"])],
        ),
        colors,
    )
    _heading(doc, "Resumo executivo", colors)
    _p(doc, answers["executive_summary"])
    _heading(doc, "Itens concluídos", colors)
    _bullets(doc, answers["completed_items"])
    _heading(doc, "Itens em andamento", colors)
    _bullets(doc, answers["in_progress_items"])
    _heading(doc, "Riscos e impedimentos", colors)
    _bullets(doc, answers["risks"])
    _heading(doc, "Próximos passos", colors)
    _bullets(doc, answers["next_steps"])
    doc.save(output_path)
    return output_path


def build_acceptance_term(scope_context: dict[str, Any], answers: dict[str, Any], output_path: Path, runtime: dict[str, Any]) -> Path:
    colors = runtime["docx"]
    identity = _identity(scope_context, answers)
    doc = Document()
    _style(doc, colors)
    _title_block(doc, "Termo de Aceite", answers["cycle_name"], identity["project_name"], colors)
    _key_value_table(
        doc,
        _metadata_rows(
            identity,
            "Termo de Aceite",
            answers["acceptance_date"],
            [
                ("Período da entrega", answers["delivery_period"]),
                ("Representante do cliente", answers["client_representative"]),
                ("Representante da contratada", answers["contractor_representative"]),
            ],
        ),
        colors,
    )
    _heading(doc, "Entregas avaliadas", colors)
    _bullets(doc, answers["delivered_items"])
    _heading(doc, "Pendências ou ressalvas", colors)
    _bullets(doc, answers["open_points"])
    _heading(doc, "Conclusão do aceite", colors)
    _p(doc, answers["final_decision"])
    _table(
        doc,
        ["Parte", "Representante", "Assinatura / Data"],
        [["Cliente", answers["client_representative"], ""], ["Contratada", answers["contractor_representative"], ""]],
        [4.5, 7.0, 5.4],
        colors,
        center_columns={0, 2},
    )
    doc.save(output_path)
    return output_path


def build_kickoff_minutes(scope_context: dict[str, Any], answers: dict[str, Any], output_path: Path, runtime: dict[str, Any]) -> Path:
    colors = runtime["docx"]
    identity = _identity(scope_context, answers)
    doc = Document()
    _style(doc, colors)
    _title_block(doc, "Ata de Reunião de Kick-off", answers["meeting_location"], identity["project_name"], colors)
    _key_value_table(
        doc,
        _metadata_rows(
            identity,
            "Ata de Reunião de Kick-off",
            answers["meeting_date"],
            [("Horário", answers["meeting_time"]), ("Local", answers["meeting_location"]), ("Objetivo da reunião", answers["meeting_objective"])],
        ),
        colors,
    )
    _heading(doc, "Participantes", colors)
    _table(
        doc,
        ["Nome", "Papel", "Organização"],
        [[row["name"], row["role"], row["organization"]] for row in answers["participants"]],
        [6.0, 5.4, 5.5],
        colors,
        center_columns={2},
    )
    _heading(doc, "Escopo inicial", colors)
    scope_rows = [[module["code"], module["title"], module["summary"]] for module in scope_context["scope_modules"]]
    _table(doc, ["Código", "Módulo", "Resumo"], scope_rows, [2.0, 5.0, 9.9], colors, center_columns={0})
    _heading(doc, "Cronograma macro", colors)
    _table(
        doc,
        ["Fase", "Período", "Objetivo"],
        [[row["phase"], row["period"], row["objective"]] for row in answers["macro_schedule"]],
        [3.0, 4.0, 9.9],
        colors,
        center_columns={0, 1},
    )
    _heading(doc, "Riscos iniciais", colors)
    _bullets(doc, answers["initial_risks"])
    _heading(doc, "Próximos passos", colors)
    _table(
        doc,
        ["Ação", "Responsável", "Prazo"],
        [[row["action"], row["owner"], row["due_date"]] for row in answers["next_steps"]],
        [8.0, 4.0, 4.9],
        colors,
        center_columns={1, 2},
    )
    doc.save(output_path)
    return output_path


def build_retrospective_minutes(scope_context: dict[str, Any], answers: dict[str, Any], output_path: Path, runtime: dict[str, Any]) -> Path:
    colors = runtime["docx"]
    identity = _identity(scope_context, answers)
    doc = Document()
    _style(doc, colors)
    _title_block(doc, "Ata de Retrospectiva", answers["sprint_name"], identity["project_name"], colors)
    _key_value_table(doc, _metadata_rows(identity, "Ata de Retrospectiva", answers["meeting_date"], [("Sprint analisada", answers["sprint_name"])]), colors)
    _heading(doc, "Participantes", colors)
    _bullets(doc, answers["participants"])
    _heading(doc, "O que deu certo", colors)
    _bullets(doc, answers["went_well"])
    _heading(doc, "O que não deu certo", colors)
    _bullets(doc, answers["didnt_go_well"])
    _heading(doc, "Pontos de melhoria", colors)
    _bullets(doc, answers["improvements"])
    _heading(doc, "Ações definidas", colors)
    _table(
        doc,
        ["Ação", "Responsável", "Prazo"],
        [[row["action"], row["owner"], row["due_date"]] for row in answers["actions"]],
        [8.0, 4.0, 4.9],
        colors,
        center_columns={1, 2},
    )
    _heading(doc, "Observações finais", colors)
    _p(doc, answers["final_notes"])
    doc.save(output_path)
    return output_path


def build_wbs_dictionary(scope_context: dict[str, Any], answers: dict[str, Any], output_path: Path, runtime: dict[str, Any]) -> Path:
    colors = runtime["docx"]
    identity = _identity(scope_context, answers)
    doc = Document()
    _style(doc, colors)
    _title_block(doc, "Dicionário da EAP", answers["dictionary_purpose"], identity["project_name"], colors)
    _key_value_table(doc, _metadata_rows(identity, "Dicionário da EAP", answers["reference_date"]), colors)
    _callout(
        doc,
        "Visão consolidada",
        "Este dicionário organiza os pacotes de trabalho derivados do escopo e ajuda a rastrear entregáveis, critérios de aceite e limites funcionais do projeto.",
        colors,
    )
    _heading(doc, "Objetivos rastreados", colors)
    _bullets(doc, scope_context["objectives"])
    _heading(doc, "Pacotes de trabalho", colors)
    _table(
        doc,
        ["Código", "Pacote", "Entregável", "Descrição", "Critério de aceite"],
        [[row["code"], row["package"], row["deliverable"], row["description"], row["acceptance"]] for row in answers["wbs_items"]],
        [1.8, 3.4, 3.6, 5.0, 3.2],
        colors,
        8.4,
        center_columns={0},
    )
    _heading(doc, "Premissas relevantes", colors)
    _bullets(doc, scope_context["assumptions"])
    _heading(doc, "Itens fora do escopo", colors)
    _bullets(doc, scope_context["out_of_scope"])
    _heading(doc, "Critérios globais de aceite", colors)
    _bullets(doc, scope_context["acceptance_criteria"])
    if scope_context["glossary"]:
        _heading(doc, "Glossário essencial", colors)
        _table(
            doc,
            ["Termo", "Definição"],
            [[row["term"], row["definition"]] for row in scope_context["glossary"]],
            [4.0, 13.0],
            colors,
            8.6,
        )
    doc.save(output_path)
    return output_path


def build_infrastructure_sizing(scope_context: dict[str, Any], answers: dict[str, Any], output_path: Path, runtime: dict[str, Any]) -> Path:
    colors = runtime["docx"]
    identity = _identity(scope_context, answers)
    doc = Document()
    _style(doc, colors)
    _title_block(doc, "Dimensionamento de Infraestrutura e Arquitetura", answers["architecture_objective"], identity["project_name"], colors)
    _key_value_table(doc, _metadata_rows(identity, "Dimensionamento de Infraestrutura e Arquitetura", answers["reference_date"]), colors)
    _callout(
        doc,
        "Escopo técnico considerado",
        "A proposta considera aplicação web responsiva, autenticação, fluxos operacionais, histórico, contato e área administrativa com separação por perfis de acesso.",
        colors,
        colors["soft_fill"],
    )
    _heading(doc, "Restrições e direcionadores", colors)
    _bullets(doc, scope_context["restrictions"])
    _heading(doc, "Ambientes previstos", colors)
    _table(
        doc,
        ["Ambiente", "Finalidade", "Disponibilidade"],
        [[row["name"], row["purpose"], row["availability"]] for row in answers["environment_rows"]],
        [3.2, 7.6, 6.2],
        colors,
        center_columns={0},
    )
    _heading(doc, "Arquitetura lógica", colors)
    _table(
        doc,
        ["Camada", "Tecnologia", "Finalidade"],
        [[row["layer"], row["technology"], row["purpose"]] for row in answers["stack_items"]],
        [4.0, 4.5, 8.0],
        colors,
    )
    _heading(doc, "Integrações e dependências", colors)
    _bullets(doc, answers["integrations"])
    _heading(doc, "Controles de segurança", colors)
    _bullets(doc, answers["security_controls"])
    _heading(doc, "Backup e recuperação", colors)
    _p(doc, answers["backup_strategy"])
    _heading(doc, "Monitoramento e observabilidade", colors)
    _p(doc, answers["monitoring_strategy"])
    _heading(doc, "Premissas operacionais", colors)
    _bullets(doc, scope_context["assumptions"])
    doc.save(output_path)
    return output_path


def build_team_directory(scope_context: dict[str, Any], answers: dict[str, Any], output_path: Path, runtime: dict[str, Any]) -> Path:
    colors = runtime["docx"]
    identity = _identity(scope_context, answers)
    doc = Document()
    _style(doc, colors)
    _title_block(doc, "Diretório da Equipe", answers["directory_purpose"], identity["project_name"], colors)
    _key_value_table(doc, _metadata_rows(identity, "Diretório da Equipe", answers["reference_date"]), colors)
    _heading(doc, "Equipe do projeto", colors)
    _table(
        doc,
        ["Nome", "Papel", "Responsabilidades", "Contato", "Disponibilidade"],
        [[row["name"], row["role"], row["responsibilities"], row["contact"], row["availability"]] for row in answers["team_directory_rows"]],
        [3.6, 3.2, 5.3, 3.0, 2.1],
        colors,
        8.4,
        center_columns={4},
    )
    _heading(doc, "Canais de comunicação", colors)
    _bullets(doc, answers["communication_channels"])
    doc.save(output_path)
    return output_path


DOCX_BUILDERS = {
    "change_request": build_change_request,
    "status_report": build_status_report,
    "acceptance_term": build_acceptance_term,
    "kickoff_minutes": build_kickoff_minutes,
    "retrospective_minutes": build_retrospective_minutes,
    "wbs_dictionary": build_wbs_dictionary,
    "infrastructure_sizing": build_infrastructure_sizing,
    "team_directory": build_team_directory,
}
