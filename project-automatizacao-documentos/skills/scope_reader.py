from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx import Document

from .utils import read_json


def _clean(text: str) -> str:
    return " ".join(text.replace("\xa0", " ").split())


def _paragraphs(doc: Document) -> list[str]:
    return [_clean(paragraph.text) for paragraph in doc.paragraphs if _clean(paragraph.text)]


def _table_rows(table) -> list[list[str]]:
    rows: list[list[str]] = []
    for row in table.rows:
        cells = [_clean(cell.text) for cell in row.cells]
        if any(cells):
            rows.append(cells)
    return rows


def _parse_sections(paragraphs: list[str]) -> dict[str, dict[str, Any]]:
    sections: dict[str, dict[str, Any]] = {}
    current_key: str | None = None
    heading_re = re.compile(r"^(\d+(?:\.\d+)*)\.?\s+(.*)$")
    for paragraph in paragraphs:
        match = heading_re.match(paragraph)
        if match:
            current_key = match.group(1)
            sections[current_key] = {"title": match.group(2), "lines": []}
            continue
        if current_key:
            sections[current_key]["lines"].append(paragraph)
    return sections


def _first_line(lines: list[str], fallback: str = "") -> str:
    return lines[0] if lines else fallback


def _infer_deliverable(title: str) -> str:
    cleaned = title.strip()
    if cleaned.lower().startswith("página"):
        return cleaned
    return f"Módulo de {cleaned.lower()} disponibilizado"


def _default_wbs_items(scope_modules: list[dict[str, Any]], acceptance_criteria: list[str]) -> list[dict[str, str]]:
    default_acceptance = acceptance_criteria[0] if acceptance_criteria else "Entrega implementada, validada e sem defeito crítico."
    rows: list[dict[str, str]] = []
    for module in scope_modules:
        rows.append(
            {
                "code": module["code"],
                "package": module["title"],
                "deliverable": _infer_deliverable(module["title"]),
                "description": module["summary"],
                "acceptance": default_acceptance,
            }
        )
    return rows


def _default_environment_rows() -> list[dict[str, str]]:
    return [
        {
            "name": "Desenvolvimento",
            "purpose": "Implementação e testes locais da equipe.",
            "availability": "Sob demanda durante o ciclo de construção.",
        },
        {
            "name": "Homologação",
            "purpose": "Validação funcional, QA e revisão com o cliente.",
            "availability": "Disponível durante as sprints de entrega e homologação.",
        },
        {
            "name": "Produção",
            "purpose": "Operação oficial do sistema após o aceite.",
            "availability": "Alta disponibilidade após o go-live.",
        },
    ]


def _default_stack_items() -> list[dict[str, str]]:
    return [
        {
            "layer": "Interface web",
            "technology": "Aplicação web responsiva",
            "purpose": "Atender clientes e administradores em navegadores atuais.",
        },
        {
            "layer": "Back-end",
            "technology": "API HTTP com regras de negócio",
            "purpose": "Autenticação, fluxos operacionais, contato, histórico e administração.",
        },
        {
            "layer": "Banco de dados",
            "technology": "Banco relacional",
            "purpose": "Persistência de usuários, registros operacionais e permissões.",
        },
        {
            "layer": "Mensageria",
            "technology": "Serviço de e-mail transacional",
            "purpose": "Recuperação de senha, confirmações e mensagens de contato.",
        },
    ]


def _default_security_controls(restrictions: list[str]) -> list[str]:
    defaults = [
        "Controle de acesso por autenticação e perfil de usuário.",
        "Rotas administrativas protegidas contra acesso indevido.",
        "Criptografia de senhas e uso de credenciais segregadas por ambiente.",
        "Registro de erros e auditoria básica para ações administrativas.",
    ]
    for item in restrictions:
        if item not in defaults:
            defaults.append(item)
    return defaults


def _default_integrations(assumptions: list[str]) -> list[str]:
    rows = [
        "Serviço de e-mail transacional para contato e recuperação de senha.",
        "Hospedagem web e banco de dados definidos pelo projeto.",
    ]
    for item in assumptions:
        if "e-mail" in item.lower() or "hosped" in item.lower():
            rows.append(item)
    return rows


def _default_communication_channels() -> list[str]:
    return [
        "Reuniões de alinhamento recorrentes com atas e decisões registradas.",
        "Canal assíncrono de mensagens para dúvidas e status rápido.",
        "Correio eletrônico para aprovações formais e envio de evidências.",
    ]


def read_scope(scope_path: Path, assets_dir: Path) -> dict[str, Any]:
    aliases = read_json(assets_dir / "scope_aliases.json")
    doc = Document(scope_path)
    paragraphs = _paragraphs(doc)
    tables = [_table_rows(table) for table in doc.tables]
    sections = _parse_sections(paragraphs)

    project: dict[str, Any] = {
        "project_name": "",
        "client_name": "",
        "contractor_name": "",
        "coordinator_name": "",
        "scope_version": "",
        "scope_date": "",
        "scope_path": str(scope_path),
    }

    if tables:
        for row in tables[0]:
            if len(row) >= 2:
                key = aliases["project_table"].get(row[0])
                if key:
                    project[key] = row[1]

    approvers: list[dict[str, str]] = []
    team: list[dict[str, str]] = []
    glossary: list[dict[str, str]] = []

    if len(tables) > 1:
        for row in tables[1][1:]:
            if len(row) >= 3:
                approvers.append({"name": row[0], "organization_role": row[1], "project_role": row[2]})
    if len(tables) > 2:
        for row in tables[2][1:]:
            if len(row) >= 3:
                team.append({"name": row[0], "role": row[1], "responsibilities": row[2]})
    if len(tables) > 3:
        for row in tables[3][1:]:
            if len(row) >= 2:
                glossary.append({"term": row[0], "definition": row[1]})

    scope_modules: list[dict[str, str]] = []
    for key, payload in sections.items():
        if key.startswith("5."):
            scope_modules.append(
                {
                    "code": key,
                    "title": payload["title"],
                    "summary": _first_line(payload["lines"]),
                    "details": payload["lines"],
                }
            )

    objective_lines = sections.get("4", {}).get("lines", [])
    objectives = [line for line in objective_lines if line != _first_line(objective_lines)]

    assumptions = sections.get("7", {}).get("lines", [])
    restrictions = sections.get("8", {}).get("lines", [])
    acceptance_criteria = sections.get("9", {}).get("lines", [])
    derived = {
        "wbs_items": _default_wbs_items(scope_modules, acceptance_criteria),
        "environment_rows": _default_environment_rows(),
        "stack_items": _default_stack_items(),
        "integrations": _default_integrations(assumptions),
        "security_controls": _default_security_controls(restrictions),
        "communication_channels": _default_communication_channels(),
        "backup_strategy": "Backup diário do banco de dados, retenção mínima semanal e procedimento de restauração validado em homologação.",
        "monitoring_strategy": "Monitoramento de disponibilidade da aplicação, filas de erro e eventos críticos de autenticação e operação.",
        "team_directory_rows": [
            {
                "name": member["name"],
                "role": member["role"],
                "responsibilities": member["responsibilities"],
                "contact": "",
                "availability": "",
            }
            for member in team
        ],
        "project_summary": _first_line(objective_lines, "Consolidar e executar o projeto conforme o escopo aprovado."),
    }

    return {
        "project": project,
        "approvers": approvers,
        "team": team,
        "glossary": glossary,
        "sections": sections,
        "objectives": objectives,
        "scope_modules": scope_modules,
        "out_of_scope": sections.get("6", {}).get("lines", []),
        "assumptions": assumptions,
        "restrictions": restrictions,
        "acceptance_criteria": acceptance_criteria,
        "derived": derived,
        "paragraphs": paragraphs,
    }
